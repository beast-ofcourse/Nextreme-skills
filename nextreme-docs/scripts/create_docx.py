#!/usr/bin/env python3
"""
nextreme-docs — spec → .docx engine

Reads a YAML or JSON spec and produces a disciplined, style-driven .docx.
Every block-level format lives in a named style. No direct-formatting headings,
no typed TOC, no unstyled tables.

Usage:
    python scripts/create_docx.py --spec templates/report_spec.yaml --output output.docx
    python scripts/create_docx.py --spec spec.json --output output.docx --pdf
    python scripts/create_docx.py --spec spec.yaml --output output.docx --doc

Spec schema (YAML or JSON):
    title: "Document Title"
    subtitle: "Optional subtitle"
    author: "Author Name"
    subject: "Subject"
    keywords: ["kw1", "kw2"]
    paper: "A4" | "Letter"                # default A4
    margins: {top: 0.85, bottom: 0.85, left: 0.85, right: 0.85}  # inches
    accent_hex: "1B4F72"                  # without #
    header_text: "CONFIDENTIAL — Title"
    footer_page_numbers: true             # default true
    toc: true                             # field-backed TOC
    different_first_page: true            # cover without header/footer
    content:
      - type: "heading"
        level: 1
        text: "Introduction"
      - type: "paragraph"
        style: "Normal"                   # or Quote, Intense Quote
        text: "Body text with **bold** and *italic* runs via markdown-ish **."
      - type: "bullet_list"
        items: ["Point one", "Point two"]
      - type: "numbered_list"
        items: ["Step one", "Step two"]
      - type: "table"
        caption: "Table 1 — Example"
        caption_position: "above"         # above | below (default below)
        style: "Light Grid Accent 1"
        headers: ["Col A", "Col B"]
        rows: [["a1", "b1"], ["a2", "b2"]]
        column_widths: [0.5, 0.5]         # proportions, sum 1.0
      - type: "image"
        path: "assets/figure.png"
        width_inches: 5.5
        caption: "Figure 1 — Caption"
        alt_text: "Description for accessibility"
      - type: "page_break"
      - type: "toc"                       # inserts TOC field at this position

Design tokens are explicit constants below — no magic numbers in logic.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    import yaml  # type: ignore
except ImportError:
    yaml = None  # type: ignore

try:
    from docx import Document
    from docx.enum.section import WD_ORIENTATION, WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Emu, Inches, Pt, RGBColor
except ImportError as import_error:
    print(f"[create_docx] Missing dependency: {import_error}", file=sys.stderr)
    print("Install: pip install \"python-docx>=0.8.11\" pyyaml", file=sys.stderr)
    sys.exit(2)


# ── Design tokens — single source of truth, no magic ────────────────────────

PAPER_A4_WIDTH = Inches(8.27)
PAPER_A4_HEIGHT = Inches(11.69)
PAPER_LETTER_WIDTH = Inches(8.5)
PAPER_LETTER_HEIGHT = Inches(11)

DEFAULT_MARGIN_TOP = Inches(0.85)
DEFAULT_MARGIN_BOTTOM = Inches(0.85)
DEFAULT_MARGIN_LEFT = Inches(0.85)
DEFAULT_MARGIN_RIGHT = Inches(0.85)
DEFAULT_HEADER_DISTANCE = Inches(0.35)
DEFAULT_FOOTER_DISTANCE = Inches(0.35)

FONT_BODY_NAME = "Calibri"
FONT_BODY_SIZE_PT = Pt(11)
FONT_HEADING_NAME = "Calibri Light"
FONT_TITLE_SIZE_PT = Pt(28)
FONT_SUBTITLE_SIZE_PT = Pt(14)
FONT_HEADING_1_SIZE_PT = Pt(16)
FONT_HEADING_2_SIZE_PT = Pt(13)
FONT_HEADING_3_SIZE_PT = Pt(11)
FONT_CAPTION_SIZE_PT = Pt(9)
FONT_HEADER_FOOTER_SIZE_PT = Pt(8)

TEXT_PRIMARY_HEX = "262626"
TEXT_SECONDARY_HEX = "595959"
TEXT_MUTED_HEX = "808080"
ACCENT_DEFAULT_HEX = "1B4F72"
ACCENT_LIGHT_HEX = "D6EAF8"
TABLE_GRID_HEX = "BDC3C7"

PARAGRAPH_SPACE_AFTER_PT = Pt(8)
PARAGRAPH_LINE_SPACING_MULTIPLE = 1.07
HEADING_1_SPACE_BEFORE_PT = Pt(18)
HEADING_1_SPACE_AFTER_PT = Pt(6)
HEADING_2_SPACE_BEFORE_PT = Pt(14)
HEADING_2_SPACE_AFTER_PT = Pt(4)
HEADING_3_SPACE_BEFORE_PT = Pt(10)
HEADING_3_SPACE_AFTER_PT = Pt(3)
CAPTION_SPACE_BEFORE_PT = Pt(4)
CAPTION_SPACE_AFTER_PT = Pt(10)

DEFAULT_TABLE_STYLE = "Light Grid Accent 1"

SLOP_TOKENS = ["lorem ipsum", "lorem", "placeholder", "your text here", "insert text here"]


# ── Data model ───────────────────────────────────────────────────────────────

@dataclass(frozen=True)
class PageSetup:
    paper: str
    margin_top: float
    margin_bottom: float
    margin_left: float
    margin_right: float
    header_distance: float
    footer_distance: float


@dataclass(frozen=True)
class DocumentConfig:
    title: str
    subtitle: str
    author: str
    subject: str
    keywords: list[str]
    accent_hex: str
    header_text: str
    footer_page_numbers: bool
    include_toc: bool
    different_first_page: bool
    page_setup: PageSetup


# ── Helpers — each does one job ─────────────────────────────────────────────

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="nextreme-docs spec → docx")
    parser.add_argument("--spec", required=True, help="Path to YAML or JSON spec")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--pdf", action="store_true", help="Also convert to PDF via LibreOffice")
    parser.add_argument("--doc", action="store_true", help="Also convert to legacy .doc via LibreOffice")
    return parser.parse_args()


def load_spec(spec_path: Path) -> dict[str, Any]:
    if not spec_path.exists():
        raise FileNotFoundError(f"Spec file not found: {spec_path}")

    raw_text = spec_path.read_text(encoding="utf-8")
    suffix = spec_path.suffix.lower()

    if suffix in (".yaml", ".yml"):
        if yaml is None:
            raise RuntimeError("pyyaml not installed — pip install pyyaml or use JSON spec")
        parsed = yaml.safe_load(raw_text)
        if not isinstance(parsed, dict):
            raise ValueError(f"YAML spec must be a mapping at top level, got {type(parsed).__name__}")
        return parsed

    if suffix == ".json":
        parsed_json = json.loads(raw_text)
        if not isinstance(parsed_json, dict):
            raise ValueError(f"JSON spec must be an object at top level, got {type(parsed_json).__name__}")
        return parsed_json

    raise ValueError(f"Unsupported spec extension '{suffix}' — use .yaml/.yml or .json")


def build_document_config(raw: dict[str, Any]) -> DocumentConfig:
    margins = raw.get("margins", {}) if isinstance(raw.get("margins"), dict) else {}

    page_setup = PageSetup(
        paper=str(raw.get("paper", "A4")),
        margin_top=float(margins.get("top", 0.85)),
        margin_bottom=float(margins.get("bottom", 0.85)),
        margin_left=float(margins.get("left", 0.85)),
        margin_right=float(margins.get("right", 0.85)),
        header_distance=float(margins.get("header_distance", 0.35)),
        footer_distance=float(margins.get("footer_distance", 0.35)),
    )

    keywords_raw = raw.get("keywords", [])
    keywords: list[str] = []
    if isinstance(keywords_raw, list):
        keywords = [str(item) for item in keywords_raw]
    elif isinstance(keywords_raw, str) and keywords_raw.strip():
        keywords = [keywords_raw.strip()]

    return DocumentConfig(
        title=str(raw.get("title", "Untitled Document")),
        subtitle=str(raw.get("subtitle", "")),
        author=str(raw.get("author", "")),
        subject=str(raw.get("subject", "")),
        keywords=keywords,
        accent_hex=str(raw.get("accent_hex", ACCENT_DEFAULT_HEX)).lstrip("#").upper() or ACCENT_DEFAULT_HEX,
        header_text=str(raw.get("header_text", "")),
        footer_page_numbers=bool(raw.get("footer_page_numbers", True)),
        include_toc=bool(raw.get("toc", False)),
        different_first_page=bool(raw.get("different_first_page", False)),
        page_setup=page_setup,
    )


def compute_content_width(section: Any) -> int:
    return int(section.page_width.emu - section.left_margin.emu - section.right_margin.emu)


def apply_page_setup(document: Any, setup: PageSetup) -> None:
    section = document.sections[0]

    paper_upper = setup.paper.upper()
    if paper_upper == "LETTER":
        section.page_width = PAPER_LETTER_WIDTH
        section.page_height = PAPER_LETTER_HEIGHT
    elif paper_upper == "A4":
        section.page_width = PAPER_A4_WIDTH
        section.page_height = PAPER_A4_HEIGHT
    else:
        raise ValueError(f"Unknown paper '{setup.paper}' — expected A4 or Letter")

    section.top_margin = Inches(setup.margin_top)
    section.bottom_margin = Inches(setup.margin_bottom)
    section.left_margin = Inches(setup.margin_left)
    section.right_margin = Inches(setup.margin_right)
    section.header_distance = Inches(setup.header_distance)
    section.footer_distance = Inches(setup.footer_distance)
    section.gutter = Inches(0)
    section.orientation = WD_ORIENTATION.PORTRAIT


def configure_styles(document: Any, accent_hex: str) -> None:
    accent_color = RGBColor.from_string(accent_hex)
    primary_color = RGBColor.from_string(TEXT_PRIMARY_HEX)
    secondary_color = RGBColor.from_string(TEXT_SECONDARY_HEX)
    muted_color = RGBColor.from_string(TEXT_MUTED_HEX)

    # Normal — body
    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY_NAME
    normal.font.size = FONT_BODY_SIZE_PT
    normal.font.color.rgb = primary_color
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = PARAGRAPH_SPACE_AFTER_PT
    normal.paragraph_format.line_spacing = PARAGRAPH_LINE_SPACING_MULTIPLE
    normal.paragraph_format.widow_control = True

    # Title
    title = document.styles["Title"]
    title.font.name = FONT_HEADING_NAME
    title.font.size = FONT_TITLE_SIZE_PT
    title.font.color.rgb = accent_color
    title.font.bold = True
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True

    # Subtitle
    subtitle = document.styles["Subtitle"]
    subtitle.font.name = FONT_HEADING_NAME
    subtitle.font.size = FONT_SUBTITLE_SIZE_PT
    subtitle.font.color.rgb = secondary_color
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)

    # Heading 1
    heading1 = document.styles["Heading 1"]
    heading1.font.name = FONT_HEADING_NAME
    heading1.font.size = FONT_HEADING_1_SIZE_PT
    heading1.font.color.rgb = accent_color
    heading1.font.bold = True
    heading1.paragraph_format.space_before = HEADING_1_SPACE_BEFORE_PT
    heading1.paragraph_format.space_after = HEADING_1_SPACE_AFTER_PT
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.keep_together = True
    heading1.paragraph_format.outline_level = 0

    # Heading 2
    heading2 = document.styles["Heading 2"]
    heading2.font.name = FONT_BODY_NAME
    heading2.font.size = FONT_HEADING_2_SIZE_PT
    heading2.font.color.rgb = accent_color
    heading2.font.bold = True
    heading2.paragraph_format.space_before = HEADING_2_SPACE_BEFORE_PT
    heading2.paragraph_format.space_after = HEADING_2_SPACE_AFTER_PT
    heading2.paragraph_format.keep_with_next = True
    heading2.paragraph_format.outline_level = 1

    # Heading 3
    heading3 = document.styles["Heading 3"]
    heading3.font.name = FONT_BODY_NAME
    heading3.font.size = FONT_HEADING_3_SIZE_PT
    heading3.font.color.rgb = accent_color
    heading3.font.bold = True
    heading3.paragraph_format.space_before = HEADING_3_SPACE_BEFORE_PT
    heading3.paragraph_format.space_after = HEADING_3_SPACE_AFTER_PT
    heading3.paragraph_format.keep_with_next = True
    heading3.paragraph_format.outline_level = 2

    # Caption
    try:
        caption = document.styles.add_style("Caption", 1)
    except ValueError:
        caption = document.styles["Caption"]
    caption.font.name = FONT_BODY_NAME
    caption.font.size = FONT_CAPTION_SIZE_PT
    caption.font.italic = True
    caption.font.color.rgb = secondary_color
    caption.paragraph_format.space_before = CAPTION_SPACE_BEFORE_PT
    caption.paragraph_format.space_after = CAPTION_SPACE_AFTER_PT
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header / Footer
    for style_name in ("Header", "Footer"):
        try:
            header_footer = document.styles.add_style(style_name, 1)
        except ValueError:
            header_footer = document.styles[style_name]
        header_footer.font.name = FONT_BODY_NAME
        header_footer.font.size = FONT_HEADER_FOOTER_SIZE_PT
        header_footer.font.color.rgb = muted_color
        header_footer.paragraph_format.space_before = Pt(0)
        header_footer.paragraph_format.space_after = Pt(0)


def configure_headers_footers(document: Any, config: DocumentConfig) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = config.different_first_page

    # Even/odd not enabled by default — user can request via spec flag if needed
    # Header
    if config.header_text:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.text = config.header_text
        paragraph.style = document.styles["Header"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if config.different_first_page:
            # First page header stays empty by intent — cover without header
            first_header = section.first_page_header
            if first_header.paragraphs:
                first_header.paragraphs[0].text = ""

    # Footer — page numbers
    if config.footer_page_numbers:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style = document.styles["Footer"]
        # Clear any default text
        paragraph.text = ""
        paragraph.add_run("Page ")
        append_field_page_number(paragraph)
        paragraph.add_run(" of ")
        append_field_num_pages(paragraph)

        if config.different_first_page:
            first_footer = section.first_page_footer
            if first_footer.paragraphs:
                first_footer.paragraphs[0].text = ""
                first_footer.paragraphs[0].style = document.styles["Footer"]


def append_field_page_number(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("w:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def append_field_num_pages(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("w:space"), "preserve")
    instr.text = " NUMPAGES "
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def append_toc_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("w:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    # Placeholder — Word replaces on update
    placeholder = paragraph.add_run()
    placeholder.text = ""

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def set_core_properties(document: Any, config: DocumentConfig) -> None:
    core = document.core_properties
    core.title = config.title
    core.subject = config.subject
    core.author = config.author
    if config.keywords:
        core.keywords = ", ".join(config.keywords)


def add_rich_paragraph(document: Any, text: str, style_name: str) -> Any:
    """
    Markdown-ish inline: **bold** and *italic* are converted to runs.
    Keeps block style intact — only inline runs vary.
    """
    paragraph = document.add_paragraph(style=style_name)

    if not text:
        return paragraph

    # Tokenize on **bold** and *italic* — non-nested, simple
    token_re = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
    parts = token_re.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

    return paragraph


def add_table_block(document: Any, block: dict[str, Any], content_width_emu: int) -> None:
    caption: str = str(block.get("caption", "")).strip()
    caption_position: str = str(block.get("caption_position", "below")).lower()
    table_style: str = str(block.get("style", DEFAULT_TABLE_STYLE))
    headers: list[str] = [str(h) for h in block.get("headers", [])]
    rows: list[list[str]] = [[str(c) for c in r] for r in block.get("rows", [])]
    column_proportions: list[float] | None = block.get("column_widths")

    if not headers and not rows:
        raise ValueError("Table block requires 'headers' or 'rows'")

    column_count = len(headers) if headers else (len(rows[0]) if rows else 0)
    if column_count == 0:
        raise ValueError("Table has zero columns — provide headers or row data")

    # Caption above
    if caption and caption_position == "above":
        paragraph = document.add_paragraph(caption, style="Caption")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=1, cols=column_count)
    table.style = table_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    header_cells = table.rows[0].cells
    if headers:
        for idx, header_text in enumerate(headers):
            cell = header_cells[idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(header_text)
            run.bold = True
            # OOXML note: header fill via w:shd is handled by table style; direct shading
            # would override theme. Rely on style's header fill for portability.
    else:
        # No headers — first data row occupies the header row slot
        for idx, cell_text in enumerate(rows[0]):
            header_cells[idx].text = str(cell_text)
        rows = rows[1:]

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for idx in range(min(len(row_data), column_count)):
            row_cells[idx].text = str(row_data[idx])

    # Column widths — proportional
    if column_proportions is not None:
        if len(column_proportions) != column_count:
            raise ValueError(
                f"column_widths length {len(column_proportions)} != column count {column_count}"
            )
        total_proportion = sum(column_proportions)
        if abs(total_proportion - 1.0) > 0.01:
            raise ValueError(f"column_widths must sum to 1.0, got {total_proportion}")
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = int(content_width_emu * column_proportions[idx])
        table.autofit = False

    # TableLook — header + banded rows
    tblPr = table._tbl.tblPr
    # Remove existing tblLook if python-docx already added one
    for existing in tblPr.findall(qn("w:tblLook")):
        tblPr.remove(existing)
    tblLook = OxmlElement("w:tblLook")
    tblLook.set(qn("w:firstRow"), "1")
    tblLook.set(qn("w:firstColumn"), "0")
    tblLook.set(qn("w:lastColumn"), "0")
    tblLook.set(qn("w:noHBand"), "0")
    tblLook.set(qn("w:noVBand"), "1")
    tblPr.append(tblLook)

    # Caption below (default)
    if caption and caption_position != "above":
        paragraph = document.add_paragraph(caption, style="Caption")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_image_block(document: Any, block: dict[str, Any], content_width_emu: int) -> None:
    image_path_raw: str = str(block.get("path", "")).strip()
    if not image_path_raw:
        raise ValueError("Image block requires 'path'")

    image_path = Path(image_path_raw)
    if not image_path.exists():
        raise FileNotFoundError(f"Image not found: {image_path}")

    caption: str = str(block.get("caption", "")).strip()
    alt_text: str = str(block.get("alt_text", caption or "Figure")).strip()
    width_inches: float | None = block.get("width_inches")
    if width_inches is not None:
        width_inches = float(width_inches)

    content_width_in = content_width_emu / 914400
    max_width_in = content_width_in - 0.1
    if width_inches is None:
        width_inches = min(5.5, max_width_in)
    if width_inches > max_width_in:
        width_inches = max_width_in

    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(width_inches))
    # Alt text — accessibility + Google Docs
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", caption or alt_text)

    if caption:
        caption_paragraph = document.add_paragraph(caption, style="Caption")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_content_blocks(document: Any, content: list[Any], content_width_emu: int) -> None:
    if not isinstance(content, list):
        raise ValueError(f"'content' must be a list, got {type(content).__name__}")

    for block in content:
        if not isinstance(block, dict):
            raise ValueError(f"Each content block must be a mapping, got {type(block).__name__}")

        block_type = str(block.get("type", "")).strip().lower()
        if not block_type:
            raise ValueError("Content block missing 'type'")

        if block_type == "heading":
            level = int(block.get("level", 1))
            if level < 1 or level > 9:
                raise ValueError(f"Heading level must be 1–9, got {level}")
            text = str(block.get("text", "")).strip()
            if not text:
                raise ValueError("Heading block requires 'text'")
            document.add_heading(text, level=level)

        elif block_type == "paragraph":
            style_name = str(block.get("style", "Normal"))
            text = str(block.get("text", ""))
            paragraph = add_rich_paragraph(document, text, style_name)
            # Optional alignment override
            alignment_raw = str(block.get("alignment", "")).lower()
            if alignment_raw == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment_raw == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment_raw == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif block_type == "bullet_list":
            items = block.get("items", [])
            if not isinstance(items, list) or not items:
                raise ValueError("bullet_list requires non-empty 'items' list")
            for item_text in items:
                document.add_paragraph(str(item_text), style="List Bullet")

        elif block_type == "numbered_list":
            items = block.get("items", [])
            if not isinstance(items, list) or not items:
                raise ValueError("numbered_list requires non-empty 'items' list")
            for item_text in items:
                document.add_paragraph(str(item_text), style="List Number")

        elif block_type == "multilevel_list":
            items = block.get("items", [])
            if not isinstance(items, list) or not items:
                raise ValueError("multilevel_list requires non-empty 'items' list (each item is {text, level})")
            # Create a single multilevel numbering definition for this block (up to 3 levels)
            numbering_id = ensure_multilevel_numbering(document, levels=[
                {"format": "decimal", "text": "%1.", "left": 0.25, "hanging": 0.25},
                {"format": "decimal", "text": "%1.%2.", "left": 0.50, "hanging": 0.35},
                {"format": "decimal", "text": "%1.%2.%3.", "left": 0.75, "hanging": 0.45},
            ])
            for raw_item in items:
                if isinstance(raw_item, dict):
                    text_value = str(raw_item.get("text", "")).strip()
                    if not text_value:
                        raise ValueError("multilevel_list item dict requires non-empty 'text'")
                    level_value = int(raw_item.get("level", 0))
                    if level_value < 0 or level_value > 2:
                        raise ValueError(f"multilevel_list level must be 0–2, got {level_value}")
                elif isinstance(raw_item, str):
                    text_value = raw_item.strip()
                    level_value = 0
                    if not text_value:
                        raise ValueError("multilevel_list string item must be non-empty")
                else:
                    raise ValueError(f"multilevel_list item must be string or mapping, got {type(raw_item).__name__}")
                paragraph = document.add_paragraph(text_value, style="Normal")
                apply_numbering(paragraph, numbering_id, level_value)

        elif block_type == "table":
            add_table_block(document, block, content_width_emu)

        elif block_type == "image":
            add_image_block(document, block, content_width_emu)

        elif block_type == "page_break":
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)

        elif block_type == "toc":
            paragraph = document.add_paragraph(style="Normal")
            append_toc_field(paragraph)

        elif block_type == "section_break":
            orientation_raw = str(block.get("orientation", "portrait")).lower()
            new_section = document.add_section(WD_SECTION.NEW_PAGE)
            if orientation_raw == "landscape":
                new_section.orientation = WD_ORIENTATION.LANDSCAPE
                new_section.page_width = PAPER_LETTER_HEIGHT if document.sections[0].page_width == PAPER_LETTER_WIDTH else PAPER_A4_HEIGHT
                new_section.page_height = PAPER_LETTER_WIDTH if document.sections[0].page_height == PAPER_LETTER_HEIGHT else PAPER_A4_WIDTH
            else:
                new_section.orientation = WD_ORIENTATION.PORTRAIT

        else:
            raise ValueError(f"Unknown block type '{block_type}' — expected heading|paragraph|bullet_list|numbered_list|multilevel_list|table|image|page_break|toc|section_break")


def check_slop_markers(text_blocks: list[str]) -> list[str]:
    findings: list[str] = []
    for block_text in text_blocks:
        lower = block_text.lower()
        for token in SLOP_TOKENS:
            if token in lower:
                findings.append(f"Slop token '{token}' in: {block_text[:80]!r}")
        if re.search(r"\bTODO\b", block_text) and "CONTENT REQUIRED" not in block_text:
            findings.append(f"Bare TODO without ticket in: {block_text[:80]!r}")
    return findings


def convert_with_libreoffice(source_path: Path, target_format: str) -> Path:
    if shutil.which("soffice") is None:
        raise RuntimeError(
            "LibreOffice 'soffice' not on PATH — install LibreOffice to convert to "
            f"'{target_format}' (https://www.libreoffice.org/download/download/)"
        )
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", target_format, str(source_path), "--outdir", str(source_path.parent)],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion to {target_format} failed: {result.stderr.strip()}")

    # soffice names output as <stem>.<format>
    expected = source_path.with_suffix(f".{target_format}")
    if not expected.exists():
        # LibreOffice may lowercase or vary — find by stem
        candidates = list(source_path.parent.glob(f"{source_path.stem}.*"))
        for candidate in candidates:
            if candidate.suffix.lower() == f".{target_format.lower()}":
                return candidate
        raise FileNotFoundError(f"LibreOffice conversion reported success but output not found near {source_path}")
    return expected


# ── Multilevel numbering — one focused helper ────────────────────────────────

def ensure_multilevel_numbering(document: Any, levels: list[dict[str, Any]]) -> int:
    """
    Create an abstractNum + num definition for multilevel lists.
    Returns the numId to apply per paragraph via apply_numbering().
    OOXML: w:abstractNum → w:num. Isolated here so generation code stays clean.
    """
    numbering_part = document.part.numbering_part
    numbering_element = numbering_part.element

    # Find next available IDs
    existing_abstract_ids = [
        int(el.get(qn("w:abstractNumId"), "0"))
        for el in numbering_element.findall(qn("w:abstractNum"))
    ]
    next_abstract_id = (max(existing_abstract_ids) + 1) if existing_abstract_ids else 0

    existing_num_ids = [int(el.get(qn("w:numId"), "0")) for el in numbering_element.findall(qn("w:num"))]
    next_num_id = (max(existing_num_ids) + 1) if existing_num_ids else 1
    # numId 0 is reserved in some viewers — start at 1 if empty

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(next_abstract_id))

    for level_index, level in enumerate(levels):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level_index))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), str(level.get("format", "decimal")))
        lvl.append(numFmt)

        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), str(level.get("text", f"%{level_index + 1}.")))
        lvl.append(lvlText)

        lvlJc = OxmlElement("w:lvlJc")
        lvlJc.set(qn("w:val"), "left")
        lvl.append(lvlJc)

        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        left_twips = int(float(level.get("left", 0.25)) * 1440)  # inches → twips
        hanging_twips = int(float(level.get("hanging", 0.25)) * 1440)
        ind.set(qn("w:left"), str(left_twips))
        ind.set(qn("w:hanging"), str(hanging_twips))
        pPr.append(ind)
        lvl.append(pPr)

        abstract_num.append(lvl)

    numbering_element.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstractNumId = OxmlElement("w:abstractNumId")
    abstractNumId.set(qn("w:val"), str(next_abstract_id))
    num.append(abstractNumId)
    numbering_element.append(num)

    return next_num_id


def apply_numbering(paragraph: Any, num_id: int, level: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


# ── Main ────────────────────────────────────────────────────────────────────

def main() -> None:
    args = parse_args()
    spec_path = Path(args.spec)
    output_path = Path(args.output)

    if output_path.suffix.lower() != ".docx":
        print(f"[create_docx] Warning: output should be .docx (got {output_path.suffix})", file=sys.stderr)

    try:
        raw_spec = load_spec(spec_path)
    except (FileNotFoundError, ValueError, RuntimeError, json.JSONDecodeError) as exc:
        print(f"[create_docx] Failed to load spec: {exc}", file=sys.stderr)
        sys.exit(1)

    try:
        config = build_document_config(raw_spec)
    except (ValueError, TypeError) as exc:
        print(f"[create_docx] Invalid config: {exc}", file=sys.stderr)
        sys.exit(1)

    # Pre-check slop in raw text fields
    text_inventory: list[str] = [config.title, config.subtitle, str(raw_spec.get("paper", ""))]
    for block in raw_spec.get("content", []):
        if isinstance(block, dict):
            for key in ("text", "caption", "header_text"):
                value = block.get(key)
                if isinstance(value, str):
                    text_inventory.append(value)
            for key in ("items", "headers", "rows"):
                value = block.get(key)
                if isinstance(value, list):
                    for item in value:
                        if isinstance(item, str):
                            text_inventory.append(item)
                        elif isinstance(item, list):
                            text_inventory.extend(str(c) for c in item)

    slop_findings = check_slop_markers(text_inventory)
    if slop_findings:
        print("[create_docx] Slop check warnings:", file=sys.stderr)
        for finding in slop_findings:
            print(f"  - {finding}", file=sys.stderr)
        print("  (Use explicit [CONTENT REQUIRED: ...] markers instead of filler)", file=sys.stderr)

    document = Document()
    try:
        apply_page_setup(document, config.page_setup)
        configure_styles(document, config.accent_hex)
        configure_headers_footers(document, config)
        set_core_properties(document, config)
    except (ValueError, RuntimeError) as exc:
        print(f"[create_docx] Setup failed: {exc}", file=sys.stderr)
        sys.exit(1)

    # Cover content
    if config.title:
        title_para = document.add_paragraph(config.title, style="Title")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if config.subtitle:
        sub_para = document.add_paragraph(config.subtitle, style="Subtitle")
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Optional top-level TOC after cover
    if config.include_toc:
        toc_heading = document.add_paragraph("Contents", style="Heading 1")
        toc_field_para = document.add_paragraph(style="Normal")
        append_toc_field(toc_field_para)

    content_width = compute_content_width(document.sections[0])
    content_blocks = raw_spec.get("content", [])

    # Allow empty content — produces a styled blank document (useful as template base)
    if content_blocks:
        try:
            add_content_blocks(document, content_blocks, content_width)
        except (ValueError, FileNotFoundError, RuntimeError) as exc:
            print(f"[create_docx] Content error: {exc}", file=sys.stderr)
            sys.exit(1)

    # Ensure parent dirs
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        document.save(str(output_path))
    except (OSError, PermissionError) as exc:
        print(f"[create_docx] Failed to save {output_path}: {exc}", file=sys.stderr)
        sys.exit(1)

    print(f"[create_docx] Wrote {output_path} ({output_path.stat().st_size} bytes)")

    # Optional conversions — always from canonical .docx
    if args.pdf:
        try:
            pdf_path = convert_with_libreoffice(output_path, "pdf")
            print(f"[create_docx] PDF: {pdf_path}")
        except (RuntimeError, FileNotFoundError) as exc:
            print(f"[create_docx] PDF conversion failed: {exc}", file=sys.stderr)
            sys.exit(1)

    if args.doc:
        try:
            doc_path = convert_with_libreoffice(output_path, "doc")
            print(f"[create_docx] DOC (OLE2 via LibreOffice): {doc_path}")
            print("[create_docx] Note: .doc is a conversion from the canonical .docx — not a native write")
        except (RuntimeError, FileNotFoundError) as exc:
            print(f"[create_docx] DOC conversion failed: {exc}", file=sys.stderr)
            sys.exit(1)


if __name__ == "__main__":
    main()
